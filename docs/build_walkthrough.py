"""Render docs/solution_walkthrough.docx from the committed artifacts.

Called by ``docs/build_docs.py``. Every number comes from ``models/*.json``, so
the walkthrough cannot drift away from the README, the web prototype, or the
artifacts they all read.
"""
from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

import config

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "solution_walkthrough.docx"

ORANGE = RGBColor(0xEB, 0x6C, 0x1E)
INK = RGBColor(0x1B, 0x1F, 0x2A)
GREY = RGBColor(0x5A, 0x61, 0x72)


def pct(x, dp=1):
    return "—" if x is None else f"{x*100:.{dp}f}%"


def num(x, dp=3):
    return "—" if x is None else f"{x:.{dp}f}"


def _base_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.12
    for name, size in (("Heading 1", 17), ("Heading 2", 13), ("Heading 3", 11.5)):
        st = doc.styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.color.rgb = ORANGE if name == "Heading 1" else INK
        st.font.bold = True
        st.paragraph_format.space_before = Pt(14 if name == "Heading 1" else 10)
        st.paragraph_format.space_after = Pt(5)


def _footer(doc: Document) -> None:
    for section in doc.sections:
        p = section.footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("AI Defense Lab for Adaptive Payment Fraud  ·  "
                        "Mastercard Innovation Challenge 2026  ·  Synthetic data only  ·  page ")
        run.font.size = Pt(8)
        run.font.color.rgb = GREY
        fld = OxmlElement("w:fldSimple")
        fld.set(qn("w:instr"), "PAGE")
        p._p.append(fld)


def _callout(doc: Document, title: str, body: str, colour=ORANGE) -> None:
    """A single-cell shaded table — a callout box that renders everywhere."""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.cell(0, 0)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "F5F6F9")
    cell._tc.get_or_add_tcPr().append(shading)
    p = cell.paragraphs[0]
    r = p.add_run(title.upper())
    r.bold = True
    r.font.size = Pt(8)
    r.font.color.rgb = colour
    p2 = cell.add_paragraph()
    r2 = p2.add_run(body)
    r2.font.size = Pt(9.5)
    doc.add_paragraph()


def _table(doc: Document, headers, rows, widths=None) -> None:
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(str(h))
        run.bold = True
        run.font.size = Pt(9)
    for r in rows:
        cells = t.add_row().cells
        for i, v in enumerate(r):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(v))
            run.font.size = Pt(9)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph()


def _figure(doc: Document, path, caption: str, width=6.4) -> None:
    if path is None or not Path(path).exists():
        return
    doc.add_picture(str(path), width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(caption)
    r.italic = True
    r.font.size = Pt(8.5)
    r.font.color.rgb = GREY


def _mono(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(8.5)
    p.paragraph_format.space_after = Pt(9)


def build_walkthrough(A: dict, figs: dict) -> Path:
    tax = A["tax"]
    counts = tax.summary_counts()
    s, m = A["summary"], A["metrics"]
    b, loao, loop = A["baseline"], A["loao"], A["loop"]
    fid, fam, ops = A["fidelity"], A["family"], A["ops"]

    from docs.build_docs import hero_family
    hf, hero = hero_family(A)

    doc = Document()
    _base_styles(doc)
    for sec in doc.sections:
        sec.left_margin = sec.right_margin = Inches(0.85)
        sec.top_margin = sec.bottom_margin = Inches(0.75)

    # ---- cover ----
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("AI DEFENSE LAB FOR ADAPTIVE PAYMENT FRAUD")
    r.bold = True
    r.font.size = Pt(26)
    r.font.color.rgb = ORANGE
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("A closed-loop GenAI red team / blue team for payment security")
    r.font.size = Pt(13)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("DISCOVER  →  SIMULATE  →  ATTACK  →  DETECT  →  ADAPT")
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = GREY
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        "Traditional fraud systems learn from fraud that has already happened. This lab "
        "actively searches for what the current defense does not know, generates realistic "
        "adversarial payment behaviour, stress-tests the model against it, and converts every "
        "discovered weakness into training data for the next defense.")
    r.italic = True
    r.font.size = Pt(11)
    doc.add_paragraph()
    _callout(doc, "Scope and honesty",
             "All data in this submission is synthetic. Every figure is a simulation result, "
             "regenerated from a single seed by `python -m src.pipeline`, and this document is "
             "generated directly from those artifacts so no number in it is typed by hand. "
             "Nothing here has been validated against real payment data, and nothing here has "
             "been reviewed or validated by Mastercard.")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Mastercard Innovation Challenge 2026 · GFF Mumbai")
    r.font.size = Pt(10)
    r.font.color.rgb = GREY
    doc.add_page_break()

    # ---- 1. executive summary ----
    doc.add_heading("1. Executive summary", level=1)
    doc.add_paragraph(
        "Generative AI has collapsed the cost of producing a new, plausible, well-targeted "
        "payment fraud attack. Personalisation at scale, fluent multilingual persuasion, "
        "synthetic voice and documents, and automated variant search mean that novel attack "
        "families now arrive faster than labelled history accumulates. A supervised fraud "
        "model trained on that history is strong on what it has seen and has no representation "
        "at all of what it has not.")
    doc.add_paragraph(
        "This project is a working laboratory for the opposite approach. It maps the emerging "
        "attack surface, has a generative red-team agent propose the next attack as a "
        "structured specification, constrains that specification against payment-domain rules, "
        "simulates it deterministically, stress-tests the defense against it, and folds "
        "whatever escapes into a cumulative replay buffer that trains the next model. Every "
        "adapted model is then a challenger that has to pass explicit promotion gates before "
        "it would go anywhere near an authorization path.")
    if hero and hf:
        _ci = hero.get("recall_after_learning_ci95")
        _ci_txt = (f", 95% interval {_ci[0]*100:.0f}-{_ci[1]*100:.0f}%" if _ci else "")
        _callout(doc, "The result that matters",
                 f"With {hf.replace('_', ' ')} removed from training entirely, the defense "
                 f"caught {pct(hero['recall_unseen'], 0)} of it. After the lab generated that "
                 f"family and replayed it into training: {pct(hero['recall_after_learning'], 0)} "
                 f"on {hero['n_test']} held-out transactions{_ci_txt}. This measures "
                 "unseen -> learned adaptation on synthetic data — not zero-shot detection, "
                 "and not production performance.")

    # ---- 2. problem ----
    doc.add_heading("2. The problem", level=1)
    doc.add_paragraph(
        "Static defenses struggle against generative-AI-enabled fraud for four compounding "
        "reasons:")
    for t, d in [
        ("Novel families have no training signal",
         "A supervised model cannot represent an attack family absent from its labels. Our "
         "leave-one-attack-family-out experiment measures exactly this collapse."),
        ("Attackers adapt faster than labels arrive",
         "Once a control starts biting, the attack changes. The label loop — fraud, dispute, "
         "chargeback, relabel, retrain — takes weeks or months."),
        ("Generation is now nearly free",
         "Producing thousands of semantically distinct lures, documents or personas costs "
         "almost nothing, which defeats template- and signature-based controls."),
        ("Much of the new fraud is authorized by the genuine customer",
         "In an authorized push payment scam every authentication signal is clean. The "
         "authorization-time model is the wrong control, and pretending otherwise produces "
         "misleading metrics."),
    ]:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(t + ". ")
        r.bold = True
        p.add_run(d)

    # ---- 3. solution ----
    doc.add_heading("3. Solution overview", level=1)
    _mono(doc, """Emerging fraud intelligence
        |
Threat Atlas             identify / threat ideation
        |
Red-team agent           structured attack specification
        |
Constraint layer         payment-domain rules: clamp or reject
        |
Constrained simulator    synthetic payment stream
        |
Defense model            authorization-time features only
     /        \\
detected    escaped
                |
       Weakness analysis      which signal did the model rely on?
                |
       Attack evolution
                |
       Adversarial replay
                |
       Retrain -> champion/challenger gate -> repeat""")
    doc.add_paragraph(
        "The two halves run on different clocks and that separation is deliberate. "
        "Authorization is synchronous and latency-bound: it reads precomputed state, scores, "
        "applies a versioned policy and logs reason codes, and it never trains. Adaptation is "
        "an offline research loop measured in days, and its output is a candidate model, not a "
        "deployment.")

    # ---- 4. identify ----
    doc.add_heading("4. Identify — the Threat Atlas", level=1)
    doc.add_paragraph(
        f"The catalog holds {counts['total_attacks']} distinct attacks across "
        f"{counts['categories']} fraud surfaces and {counts['rails']} payment rails. Each entry "
        "records what a defender can observe: the attacker's objective, the specific role "
        "generative AI plays, the kill chain, the transaction and behavioural signatures, how "
        "visible the attack is at authorization time, and which signals only appear after "
        "settlement. It contains no operational guidance of any kind.")
    _table(doc, ["Simulator status", "Count", "What it means"], [
        ("IMPLEMENTED", counts["implemented"],
         "a dedicated injector reproduces its authorization footprint"),
        ("PARAMETERIZED", counts["parameterized"],
         "reachable by configuring an existing injector, no new code"),
        ("RESEARCH_ONLY", counts["research_only"], "characterized but NOT simulated"),
        ("FUTURE", counts["future"], "named as planned simulator work"),
    ], widths=[1.7, 0.8, 3.9])
    _table(doc, ["Fraud surface", "Attacks", "Simulated"],
           [(c, d["total"], d["IMPLEMENTED"] + d["PARAMETERIZED"])
            for c, d in sorted(tax.coverage_by_category().items(),
                               key=lambda kv: -kv[1]["total"])],
           widths=[3.6, 1.0, 1.2])
    _figure(doc, figs.get("coverage"),
            "The catalog is deliberately wider than the simulator, and every entry is labelled "
            "so the gap is visible rather than implied.")
    _callout(doc, "Why the gap is stated, not hidden",
             f"{counts['auth_time_hard']} of the {counts['total_attacks']} catalogued attacks "
             "have low or no visibility at authorization time. An authorization-time model is "
             "the wrong control for most of them, and a submission that quietly counted them "
             "as covered would be overstating what it built.")

    # ---- 5. generate ----
    doc.add_heading("5. Generate — the constrained simulator", level=1)
    doc.add_heading("5.1 Where generative AI actually contributes", level=2)
    doc.add_paragraph(
        "The language model never writes transaction rows. It writes a specification — which "
        "behavioural dial to move, in which direction, and which detector signal that is meant "
        "to defeat:")
    _mono(doc, json.dumps({
        "attack_family": "account_takeover",
        "strategy": "reuse the victim's trusted device, change merchant behaviour",
        "amount_profile": "moderate", "velocity_profile": "low_and_slow",
        "device_behavior": "trusted_device", "merchant_behavior": "new_high_risk_merchant",
        "geo_behavior": "plausible", "targets_signal": "device_changed",
        "confidence": 0.82}, indent=2))
    doc.add_paragraph(
        "That specification passes through a payment-domain constraint layer before anything is "
        "generated. An authorized push payment scam cannot run from an attacker's device, "
        "because the genuine customer is the one authenticating; a probing campaign cannot be "
        "high-value, because it would not be economic. Values outside their permitted range are "
        "clamped, contradictions are corrected to the nearest legal value, and in strict mode "
        "the specification is refused outright. A deterministic, seeded simulator then executes "
        "what survives.")
    _callout(doc, "The design claim",
             "LLM creativity constrained by payment-domain simulation rules. Both paths emit "
             "the same structure and both pass through the same constraint layer, which is "
             "what lets the system be generative and reproducible at once.")
    _callout(doc, "Two paths, and which one produced the committed artifacts",
             "DEMO MODE USES DETERMINISTIC, COMMITTED SPECIFICATIONS. Every specification in "
             "models/attack_lineage.json carries spec_source \"heuristic\" and every content "
             "artifact carries source \"template\". That is a reliability decision: the "
             "prototype must run with no API key, no network and no training, and every "
             "committed figure must reproduce from a single seed. The weakness-driven "
             "heuristic reads the same measured evidence and moves the same dials, so the "
             "behaviour demonstrated is the loop's rather than one model call's. "
             "THE OPTIONAL GENAI RED TEAM GENERATES THE SPECIFICATION INSTEAD when an API key "
             "is present: the language model receives the measured weakness and returns the "
             "structured specification, which passes through exactly the same payment-domain "
             "constraint layer before anything is simulated, stamped spec_source \"llm\" so "
             "the two are never confused. Run `python -m src.generate.demo_specs` to produce "
             "that evidence; with no key it writes nothing rather than substituting heuristic "
             "output for a model response.",
             colour=GREY)

    doc.add_heading("5.2 Simulating a payment portfolio worth training on", level=2)
    if s:
        arch = s.get("customer_archetypes", {})
        doc.add_paragraph(
            f"The committed dataset holds {s['n_transactions']:,} transactions across "
            f"{s['n_cardholders']:,} cardholders and {s['n_merchants']:,} merchants, with "
            f"{s['n_fraud']} fraudulent authorizations at a {s['fraud_rate']*100:.1f}% base "
            f"rate, plus {s.get('n_cover_transactions', 0)} transactions of fraud-actor cover "
            "traffic.")
        if arch:
            _table(doc, ["Customer archetype", "Cardholders"],
                   [(k.replace("_", " "), v) for k, v in
                    sorted(arch.items(), key=lambda kv: -kv[1])], widths=[3.2, 1.4])
    doc.add_paragraph(
        "Two design decisions do most of the work in making the data honest. First, fraud "
        "actors have a history: mule accounts, bust-out accounts and front merchants build "
        "ordinary-looking traffic before they are used, and that traffic is labelled "
        "legitimate, because at authorization time it is. Without it, 'this card has no "
        "history' becomes a synonym for fraud and the model learns the generator. Second, "
        "attacks reuse things: probing revisits the same cards and merchants, mimicry shops at "
        "the victim's own regular merchant, laundering pushes many cards through one front. "
        "Without reuse every fraudulent row would trivially be a first-ever card/merchant pair.")

    doc.add_heading("5.3 Fidelity diagnostics", level=2)
    if fid:
        doc.add_paragraph(
            "If any single field separated fraud from genuine traffic cleanly, the detector "
            "would be reading the generator rather than the fraud, and every downstream number "
            "would be meaningless. So it is measured:")
        _table(doc, ["Check", "Result", "Detail"],
               [(c["check"].replace("_", " "), c["status"], c["detail"])
                for c in fid["checks"]], widths=[2.0, 0.7, 3.7])
        doc.add_paragraph(
            f"Strongest single feature: {fid['separability'][0]['feature']} at AUC "
            f"{fid['max_single_feature_auc']:.3f}. {fid['n_fail']} failing checks, "
            f"{fid['n_warn']} warnings.")
        _figure(doc, figs.get("separability"),
                "No single authorization-time feature comes close to separating the classes.")
        _callout(doc, "Scope of these diagnostics",
                 "These measure the simulator against itself. They say nothing about similarity "
                 "to any real payment portfolio, and no comparison against production data of "
                 "any kind has been performed.")

    # ---- 6. defend ----
    doc.add_heading("6. Defend — the detection stack", level=1)
    doc.add_paragraph(
        "Gradient boosting fused with an isolation forest fitted on genuine traffic only, over "
        "authorization-time features exclusively. Post-outcome fields — refunds, chargebacks, "
        "the authorization result itself — are hard-blocked in code and the block is "
        "unit-tested. Scores are isotonically calibrated on a held-out slice, so the number a "
        "reviewer sees is an estimated fraud probability rather than an arbitrary blend, and "
        "the operating threshold is chosen against an agreed false-positive budget.")
    if m:
        _table(doc, ["Metric", "Value"], [
            ("Recall", pct(m["recall"])),
            ("Precision", num(m["precision"])),
            ("F1", num(m["f1"])),
            ("PR-AUC", num(m["pr_auc"])),
            ("ROC-AUC", num(m["roc_auc"])),
            ("False-positive rate", f"{pct(m['false_positive_rate'], 2)} "
                                    f"(budget ≤ {pct(config.TARGET_MAX_FPR, 0)})"),
            ("False positives per 1,000 genuine payments",
             f"{m['false_positive_rate']*1000:.1f}"),
            ("Held-out test set", f"{m.get('test_size', 0):,} transactions, "
                                  f"{m.get('n_fraud_eval', 0)} fraudulent"),
        ], widths=[3.6, 2.8])
        doc.add_paragraph(
            f"Precision is naturally modest at a {config.DEFAULT_FRAUD_RATE*100:.1f}% base "
            "rate. The answer is tiered decisioning — challenge, do not decline — rather than "
            "a bigger number.")
    _figure(doc, figs.get("sweep"),
            "There is no magic threshold. Every point is a different answer to how much genuine "
            "friction is worth buying detection with.")

    doc.add_heading("6.1 Recall by attack family", level=2)
    if fam and "static_ml" in fam.get("models", {}):
        doc.add_paragraph(
            f"Measured on an unseen, fraud-enriched frame ({fam['frame']['n_transactions']:,} "
            f"transactions, {fam['frame']['n_fraud']} fraudulent) generated from a seed no "
            "model was trained or tuned on, so each family has enough held-out examples for its "
            "number to mean something.")
        from docs.build_docs import ci_text
        block = fam["models"]["static_ml"]["per_family_recall"]
        _table(doc, ["Attack family", "Recall", "n", "95% interval"],
               [(k, pct(v["recall"], 0), v["n"],
                 ci_text(v["recall"], v["n"], v.get("ci95")))
                for k, v in sorted(block.items(), key=lambda kv: -kv[1]["recall"])],
               widths=[2.5, 1.0, 0.7, 1.4])
        _figure(doc, figs.get("family"),
                "Recall by family with 95% Wilson intervals. The intervals are shown because "
                "quoting a proportion without its sample size is how efficacy claims fall over "
                "under questioning.")
    _callout(doc, "Authorization-time-hard fraud",
             "Families whose fraud is authorized by the genuine customer on their own device — "
             "first-party dispute abuse, and scams the victim was talked into — sit at the "
             "bottom by construction. There is very little to see at authorization time. The "
             "correct control is friction, payee-risk intelligence and post-transaction "
             "recall, not a hard decline. They are reported rather than quietly excluded.")

    doc.add_heading("6.2 Decision policy and reason codes", level=2)
    doc.add_paragraph(
        "The model does not emit a block. It emits a calibrated probability, which a "
        "prototype issuer-style policy layer maps to APPROVE, STEP-UP or DECLINE. The step-up "
        "boundary is the operating threshold already tuned against the false-positive budget, "
        "rather than a second unjustified constant. Each decision carries short reason codes "
        "derived from which rule-style signals fired — unusual device behaviour, abnormal "
        "velocity, atypical amount, new payee, geographic deviation, device shared across "
        "cards, merchant seeing almost only new cards — so an analyst sees why.")
    if ops and "static_ml" in ops:
        o = ops["static_ml"]
        sc = o["scenario"]
        _table(doc, ["Operational measure", "Value"], [
            ("False positives per 1,000 genuine payments",
             f"{o['per_1000']['false_positives_per_1000_legit']:.1f}"),
            ("Genuine customers sent to step-up",
             pct(o["action_distribution"]["genuine_customers_stepped_up"], 2)),
            ("Genuine customers declined",
             pct(o["action_distribution"]["genuine_customers_declined"], 3)),
            ("Fraud sent to friction", pct(o["policy"]["fraud_to_friction"], 0)),
            ("Monthly step-up challenges", f"{sc['monthly_step_up_challenges']:,}"),
            ("Analyst-hours if all reviewed manually",
             f"{sc['monthly_review_hours_if_all_manually_reviewed']:,}"),
        ], widths=[3.8, 2.6])
        doc.add_paragraph(f"Scenario: {sc['label']}. {o['disclaimer']}")

    # ---- 7. closed loop ----
    doc.add_heading("7. The closed loop — red team meets blue team", level=1)
    doc.add_paragraph("One round is a full red-team experiment, not a retraining step:")
    for i, t in enumerate([
        "Evaluate the defense currently in force on held-out traffic.",
        "Locate its weakest families, and measure which signals its ranking of each one "
        "actually depends on, by permuting features on that family's own rows.",
        "Propose the next attack generation as a structured specification aimed at removing "
        "that dependency.",
        "Constrain the specification against payment-domain rules.",
        "Simulate the evolved generation deterministically.",
        "Stress-test the stale defense against it — this is the gap the red team just opened.",
        "Replay what escaped into a bounded, stratified buffer holding every prior generation.",
        "Retrain a candidate on the base data plus the whole buffer.",
        "Re-evaluate on the new generation, every prior generation, families never attacked, "
        "and genuine traffic.",
        "Put the candidate through champion/challenger gates.",
        "Carry the residual weakness into the next round.",
    ], start=1):
        doc.add_paragraph(f"{i}. {t}", style="List Number")
    _callout(doc, "Why this is not a scripted stealth dial",
             "The family attacked and the dial moved are both derived from measurement. The "
             "loop attacks whichever family the current defense handles worst, and it moves "
             "whichever dial removes the signal that model was measured to lean on. Change the "
             "model and the loop attacks something else.")
    retired = (loop or {}).get("retired_frontiers", [])
    if retired:
        _callout(doc, "The red team reallocates",
                 " ".join(f"After round {r['round']} it retired {r['family']}"
                          + (f" and moved to {r['replaced_by']}." if r["replaced_by"] else ".")
                          for r in retired)
                 + " A family that stays a residual frontier under repeated attack will not "
                   "yield to more of the same. Continuing to hammer it burns replay capacity "
                   "on examples the model cannot separate, and measurably drags down families "
                   "it could. The frontier is reported, not hidden.",
                 colour=GREY)

    if loop:
        rows = []
        for h in loop["history"]:
            for famname, d in h["families"].items():
                rows.append((h["round"], famname, d.get("targets_signal") or "—",
                             pct(d["stale_recall"], 0), pct(d["adapted_recall"], 0),
                             d["status"], d["n"]))
        _table(doc, ["Round", "Family", "Signal targeted", "Stale", "After replay",
                     "Status", "n"], rows,
               widths=[0.6, 1.5, 1.6, 0.7, 0.9, 1.1, 0.5])
        _callout(doc, "Which path produced this lineage",
                 "Every node in the table above carries spec_source \"heuristic\". Demo mode "
                 "uses the deterministic committed specifications, so these numbers reproduce "
                 "from the seed with no API key. The optional GenAI red team produces the same "
                 "structure through the same constraint layer, stamped spec_source \"llm\".",
                 colour=GREY)
        _figure(doc, figs.get("lineage"),
                "Attack lineage. Each generation targets the signal the defense was measured "
                "to rely on; the dashed line is the stale defense meeting it for the first "
                "time, the solid line is the same defense after adversarial replay.")

    doc.add_heading("7.1 Champion / challenger governance", level=2)
    doc.add_paragraph(
        "An adaptive defense that deploys itself is not deployable. Every model the loop "
        "produces is a challenger measured against the model in force, and it is promoted only "
        "if it clears every gate:")
    g = config.CHAMPION_CHALLENGER
    _table(doc, ["Gate", "Requirement"], [
        ("Attack recall gain", f"≥ champion + {g['min_attack_recall_gain']:.0%}"),
        ("Absolute false-positive ceiling", f"≤ {g['max_fpr']:.1%} of genuine traffic"),
        ("False-positive regression", f"≤ champion + {g['max_fpr_regression']:.1%}"),
        ("No catastrophic forgetting",
         f"no prior family drops more than {g['max_prior_recall_drop']:.0%}"),
        ("Overall ranking quality", f"PR-AUC ≥ champion − {g['max_overall_pr_auc_drop']:.2f}"),
    ], widths=[2.6, 3.8])
    if loop:
        promo = loop.get("promotion", {})
        promoted = promo.get("promoted_rounds", [])
        if promoted:
            doc.add_paragraph(f"In the committed run, round(s) {promoted} were promoted.")
        else:
            _callout(doc, "A result we did not tidy away",
                     "In the committed run no candidate cleared every gate. The adapted models "
                     "improved detection on the evolved attack and breached another constraint "
                     "the gates protect, so they would not be deployed. Both the adapted "
                     "candidate and the promoted champion are shipped as separate artifacts, "
                     "and the benchmark tables label which is which.",
                     colour=RGBColor(0xE5, 0x48, 0x4D))

    # ---- 8. experiments ----
    doc.add_heading("8. Key experiments", level=1)
    doc.add_heading("8.1 Rules vs static ML vs adaptive defense", level=2)
    doc.add_paragraph(
        "The rule set is deliberately competent rather than a strawman: it includes the "
        "network-level rules a real fraud team would write once it had the same counters the "
        "model gets. Its thresholds are round, domain-chosen numbers, never fitted to this "
        "dataset. Detectors are compared at a matched false-positive budget, because comparing "
        "them at whatever operating point each happens to sit on is close to meaningless.")
    if b and "_fpr_matched" in b:
        labels = {"rules_baseline": "Rules baseline", "static_ml": "Static ML",
                  "adaptive_ml": "Adaptive (promoted)",
                  "adaptive_candidate_unpromoted": "Adaptive candidate (not promoted)"}
        _table(doc, ["Detector", "Recall", "Precision", "F1", "FP / 1,000 genuine"],
               [(labels.get(k, k), pct(v["recall"], 0), num(v["precision"], 3),
                 num(v["f1"], 3), f"{v.get('false_positives_per_1000_legit', 0):.1f}")
                for k, v in b["_fpr_matched"]["models"].items()],
               widths=[2.4, 0.9, 1.0, 0.8, 1.3])
        _figure(doc, figs.get("baseline"),
                "Every detector re-thresholded to spend the same false-positive budget on the "
                "same held-out split.")

    doc.add_heading("8.2 What adaptation buys — the evolved attacks", level=2)
    h2h = A.get("h2h")
    if h2h:
        doc.add_paragraph(
            "Section 8.1 scores every detector on the original attack distribution, which is "
            "the static model's home ground: it was trained on exactly that. The question "
            "adaptation exists to answer is what happens once the attack has moved. Both "
            "models below are scored on the same unseen frame of the FINAL evolved "
            f"generation ({h2h['frame']['n_fraud']} fraudulent transactions, from a seed "
            "neither model has seen.)")
        labels = {"static_defense": "Static defense (never saw the evolved attack)",
                  "adaptive_defense": "Adaptive defense (trained through the loop)",
                  "promoted_champion": "Promoted champion"}
        _table(doc, ["Defense", "Mean recall on evolved attacks", "FPR"],
               [(labels.get(k, k), pct(v["mean_evolved_recall"], 0),
                 pct(v["false_positive_rate"], 2))
                for k, v in h2h["models"].items()], widths=[3.4, 1.9, 1.1])
        rows = []
        for fam in h2h["focus"]:
            row = [fam]
            for k in h2h["models"]:
                d = h2h["models"][k]["evolved_family_recall"].get(fam)
                row.append(f"{d['recall']*100:.0f}% (n={d['n']})" if d else "—")
            rows.append(tuple(row))
        _table(doc, ["Evolved family"] + [labels.get(k, k).split(" (")[0]
                                          for k in h2h["models"]], rows)
        _callout(doc, "Why the static model can still look better elsewhere",
                 "On the original distribution the static model is at or near the top, and it "
                 "should be — that is the data it was trained on. The adaptive model carries "
                 "several generations of evolved attacks in its training set, which costs it "
                 "a little on the original distribution and buys it the numbers above. "
                 "Presenting only whichever table flattered us would be the easy version of "
                 "this submission.")
    else:
        doc.add_paragraph("Not available: run the closed loop and the diagnostics step.")

    doc.add_heading("8.3 Leave-one-attack-family-out", level=2)
    doc.add_paragraph(
        "Each family is removed from training entirely — every example, not a held-out sample "
        "— then scored as a genuinely unseen attack, then re-added. A large gain means the "
        "family is unseen-hard but learnable, which is precisely the gap the adversarial replay "
        "loop exists to close.")
    if loao:
        from docs.build_docs import MIN_N_TO_HEADLINE, ci_text
        rows, thin = [], []
        for f, r in sorted(loao["families"].items(), key=lambda kv: -(kv[1]["gain"] or 0)):
            if r["recall_unseen"] is None:
                continue
            small = r["n_test"] < MIN_N_TO_HEADLINE
            if small:
                thin.append(f)
            rows.append((f + (" *" if small else ""),
                         pct(r["recall_unseen"], 0),
                         pct(r["recall_after_learning"], 0),
                         ci_text(r["recall_after_learning"], r["n_test"],
                                 r.get("recall_after_learning_ci95")),
                         f"{r['gain']*100:+.0f} pts", r["n_test"]))
        _table(doc, ["Family", "Never seen", "After learning", "95% interval", "Gain", "n"],
               rows, widths=[1.9, 1.0, 1.1, 1.2, 0.7, 0.5])
        _figure(doc, figs.get("loao"), "Unseen versus learned, by family.")
        if thin:
            _callout(doc, "Marked * — too thin to headline",
                     f"{', '.join(thin)} rest on fewer than {MIN_N_TO_HEADLINE} held-out "
                     "examples. They are shown for completeness and deliberately not used as "
                     "the hero result: the interval is far too wide to carry a headline, "
                     "however good the point estimate looks. The hero result is selected "
                     "automatically from the families that clear the sample-size floor.",
                     colour=GREY)
        _callout(doc, "How to read this",
                 loao.get("what_this_does_not_show", ""),
                 colour=RGBColor(0xE5, 0x48, 0x4D))

    # ---- 9. hero ----
    doc.add_heading("9. Hero result", level=1)
    if hero and hf:
        doc.add_paragraph(
            f"Attack family: {hf.replace('_', ' ')}, removed from training entirely.")
        from docs.build_docs import MIN_N_TO_HEADLINE, ci_text
        _table(doc, ["Stage", "Recall", "95% interval", "n"], [
            ("Never seen — the blind spot", pct(hero["recall_unseen"], 0),
             ci_text(hero["recall_unseen"], hero["n_test"],
                     hero.get("recall_unseen_ci95")), hero["n_test"]),
            ("After the lab generated and replayed it",
             pct(hero["recall_after_learning"], 0),
             ci_text(hero["recall_after_learning"], hero["n_test"],
                     hero.get("recall_after_learning_ci95")), hero["n_test"]),
        ], widths=[2.9, 1.0, 1.4, 0.6])
        doc.add_paragraph(
            "Labelled explicitly as unseen → learned. It is not evidence of zero-shot "
            "detection, and it is not a production result.")
        doc.add_paragraph(
            f"This family is chosen automatically as the largest measured gain among families "
            f"with at least {MIN_N_TO_HEADLINE} held-out examples. Families below that floor "
            "produce wider intervals than a headline can carry, so they are reported in the "
            "table above and never led with — however good the point estimate looks.")

    # ---- 10. residual frontiers ----
    doc.add_heading("10. Residual frontiers", level=1)
    doc.add_paragraph(
        "Some attacks remain difficult after adaptation, and the loop is designed to surface "
        "them rather than smooth them over:")
    for t, d in [
        ("Behavioural mimicry on the legitimate centroid",
         "when every individual signal is held inside the cardholder's own normal range, "
         "there is very little left for an authorization-time model to separate."),
        ("Victim-authorized payments",
         "the genuine customer authenticates and authorizes. The transaction is legitimate by "
         "every signal the authorization carries."),
        ("First-party dispute abuse",
         "indistinguishable from a genuine purchase at authorization; the fraud is the later "
         "dispute, which is post-outcome by definition."),
    ]:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(t + " — ")
        r.bold = True
        p.add_run(d)
    doc.add_paragraph(
        "These are reported as residual frontiers with their measured recall. They are exactly "
        "why continuous red teaming is required, and exactly why the right control for them is "
        "friction and post-transaction intervention rather than a hard decline.")

    # ---- 11. feasibility ----
    doc.add_heading("11. Real-world feasibility", level=1)
    doc.add_paragraph(
        "Online, the authorization path reads precomputed per-card and network counters, "
        "scores, calibrates, applies a versioned policy and logs reason codes. It never "
        "trains and it never calls a language model. The relational features are running "
        "counters keyed on card, device, IP and merchant — constant-time lookups, not graph "
        "queries — and they are the signals an issuer or a single merchant cannot compute "
        "alone but a payment network can.")
    doc.add_paragraph(
        "Offline, the red-team loop generates attacks, stress-tests, replays and retrains on a "
        "research cadence, and its output passes through the promotion gates above before it "
        "could reach production. A real deployment would additionally run the challenger in "
        "shadow against live traffic, monitor drift, and feed analyst outcomes back — none of "
        "which is simulated here.")

    # ---- 12. limitations ----
    doc.add_heading("12. Limitations", level=1)
    for t in [
        "Synthetic data only. No real cardholder data is used anywhere, and every number in "
        "this document is a simulation result.",
        "No validation against real payment data. A labelled backtest on issuer data is the "
        "necessary next step before any of these figures means anything operationally.",
        "Nothing here has been reviewed or validated by Mastercard, and no part of this "
        "describes any real payment-network system.",
        "Some fraud is not observable at authorization time; those families are reported with "
        "low recall rather than excluded.",
        "Residual frontiers are not solved.",
        "The text arm is a trivially separable synthetic sanity check, never evidence of "
        "detection efficacy. Its corpus is composed from a fixed slot vocabulary, so the two "
        "classes separate on vocabulary alone and the score is high for that reason and no "
        "other. It measures the corpus, not the detector, and would not survive contact with "
        "real scam messages. Every detection claim in this submission rests on the transaction "
        "model alone.",
        "Network features are computed over synthetic relationships only.",
        "No shadow rollout, drift monitoring, analyst feedback loop, or production latency and "
        "availability engineering.",
    ]:
        doc.add_paragraph(t, style="List Bullet")

    # ---- 13. future work ----
    doc.add_heading("13. Future work", level=1)
    for t in [
        "Validation against a labelled issuer backtest, and a cross-check on a public fraud "
        "benchmark.",
        "UPI and account-to-account as first-class simulated rails, with VPA-style payee "
        "identity and instant-settlement semantics.",
        "Richer entity-graph features for coordinated rings, beyond the current counter-based "
        "relational signals.",
        "Oracle-guided evasion, where the red team probes the score surface directly rather "
        "than reasoning from measured feature attribution.",
        "Multilingual and code-mixed scam-text generation, with detection reported separately "
        "on the non-English slice.",
        "Shadow-mode deployment simulation and drift monitoring to complete the governance "
        "story.",
    ]:
        doc.add_paragraph(t, style="List Bullet")

    doc.add_heading("14. Reproducing every number in this document", level=1)
    _mono(doc, "python -m pip install -r requirements.txt\n"
               "python -m src.pipeline          # regenerate every artifact\n"
               "python -m pytest tests/ -q      # test suite\n"
               "python -m docs.build_docs       # regenerate README + this document\n"
               "streamlit run app/Home.py       # the web prototype")
    doc.add_paragraph(
        f"Everything flows from a single seed (config.GLOBAL_SEED = {config.GLOBAL_SEED}, "
        f"schema version {config.SCHEMA_VERSION}). This document is generated directly from "
        "the committed artifacts, so it cannot disagree with the README or the web prototype: "
        "they all read the same files.")

    props = doc.core_properties
    props.title = "AI Defense Lab for Adaptive Payment Fraud"
    props.subject = "Closed-loop GenAI red team / blue team for payment security"
    props.category = "Mastercard Innovation Challenge 2026"
    props.comments = ("Generated from committed artifacts by docs/build_docs.py. "
                      "All data synthetic.")

    _footer(doc)
    doc.save(OUT)
    return OUT
